import os
from strands import Agent, tool
from strands.models import BedrockModel
from docx import Document
from pypdf import PdfReader

from agents.config.config import input_folder_dir_path, model_id_claude3_7, model_temperature


# Create a BedrockModel
bedrock_model = BedrockModel(
    model_id=model_id_claude3_7,
    temperature=model_temperature
)

def read_file_from_input_dir(filename):
    """Read file from the case-specific input directory"""
    from agents.utils.project_context import get_case_input_directory
    input_dir = get_case_input_directory()
    full_path = os.path.join(input_dir, filename)
    return full_path

def find_mra_file():
    """Find the MRA file in the input directory (supports .md, .docx, .pdf)"""
    from agents.utils.project_context import get_case_input_directory
    input_dir = get_case_input_directory()
    
    # Check for MRA file with any supported extension
    mra_patterns = [
        'mra-assessment.pdf',
        'mra-assessment.docx', 
        'mra-assessment.md',
        'mra-assessment.doc',
        # Legacy patterns for backward compatibility
        'aws-customer-migration-readiness-assessment.md',
        'customer-assessment-summary.pdf'
    ]
    
    for pattern in mra_patterns:
        filepath = os.path.join(input_dir, pattern)
        if os.path.exists(filepath):
            return pattern
    
    # If no file found, list what's available for debugging
    try:
        files = os.listdir(input_dir)
        print(f"Available files in input directory: {files}")
    except Exception as e:
        print(f"Error listing input directory: {e}")
    
    return None

@tool(name="read_docx_file", description="Read Word document (.docx) from the input folder and extract text content")
def read_docx_file(filename: str):
    """Read Word document and extract text content"""
    full_path = read_file_from_input_dir(filename)
    doc = Document(full_path)
    content = []
    
    for para_num, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            content.append(para.text)
    
    # Extract tables if any
    for table_num, table in enumerate(doc.tables, 1):
        content.append(f"\n--- Table {table_num} ---")
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            content.append(row_text)
    
    return "\n".join(content)

@tool(name="read_markdown_file", description="Read Markdown file (.md) from the input folder and extract text content")
def read_markdown_file(filename: str):
    """Read Markdown file and extract text content"""
    full_path = read_file_from_input_dir(filename)
    with open(full_path, 'r', encoding='utf-8') as file:
        content = file.read()
    return content

@tool(name="read_pdf_file", description="Read PDF file (.pdf) from the input folder and extract text content")
def read_pdf_file(filename: str):
    """Read PDF file and extract text content"""
    full_path = read_file_from_input_dir(filename)
    reader = PdfReader(full_path)
    content = []
    
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text.strip():
            content.append(f"--- Page {page_num} ---\n{text}")
    
    return "\n\n".join(content)

# System message for the MRA analysis agent
system_message = """
You are an AWS Migration Readiness Assessment (MRA) specialist with expertise in evaluating 
organizational readiness for cloud migration and transformation.

**About MRA**: Migration Readiness Assessment is a comprehensive evaluation framework that assesses 
an organization's preparedness across multiple dimensions to successfully migrate to AWS. It identifies 
gaps, risks, and provides actionable recommendations to improve migration readiness.

You have access to tools to read MRA documents in various formats:
- **read_docx_file**: Read Word documents (.docx, .doc)
- **read_markdown_file**: Read Markdown files (.md)
- **read_pdf_file**: Read PDF files (.pdf)

**IMPORTANT**: The MRA file will be named 'mra-assessment' with the appropriate extension (.pdf, .docx, or .md).
First, try reading 'mra-assessment.pdf', then 'mra-assessment.docx', then 'mra-assessment.md' until you find the file.
If none of these work, try legacy filenames like 'aws-customer-migration-readiness-assessment.md' or 'customer-assessment-summary.pdf'.

When analyzing MRA documents, focus on extracting and synthesizing:

## (1) Executive Summary & Assessment Overview
- Overall migration readiness score/maturity level
- Key findings and critical observations
- Assessment methodology and scope
- Stakeholders involved and their roles

## (2) Business Readiness
- **Business Case & Strategy**:
  - Strategic alignment with business objectives
  - Executive sponsorship and commitment
  - Business case clarity and ROI expectations
  - Change management readiness
- **Organizational Structure**:
  - Cloud Center of Excellence (CCoE) maturity
  - Roles and responsibilities definition
  - Decision-making processes
  - Governance framework

## (3) People & Culture Readiness
- **Skills & Capabilities**:
  - Current cloud skills inventory
  - Skill gaps and training needs
  - Hiring and retention strategies
  - Partner ecosystem engagement
- **Culture & Change Management**:
  - Organizational change readiness
  - Innovation culture and mindset
  - Risk tolerance and appetite
  - Communication and engagement plans

## (4) Process Readiness
- **Migration Process Maturity**:
  - Migration methodology and approach
  - Wave planning and prioritization
  - Testing and validation processes
  - Cutover and rollback procedures
- **Operational Processes**:
  - ITSM and ITIL process maturity
  - Incident and problem management
  - Change and release management
  - Service level management

## (5) Technology & Platform Readiness
- **Current State Technology**:
  - Application portfolio assessment
  - Infrastructure and architecture review
  - Technical debt and dependencies
  - Integration complexity
- **AWS Platform Readiness**:
  - Landing zone design and implementation
  - Network connectivity and architecture
  - Security and compliance controls
  - Monitoring and observability setup

## (6) Security & Compliance Readiness
- **Security Posture**:
  - Security framework and policies
  - Identity and access management
  - Data protection and encryption
  - Threat detection and response
- **Compliance Requirements**:
  - Regulatory compliance needs
  - Industry standards adherence
  - Audit and reporting capabilities
  - Data residency and sovereignty

## (7) Operations Readiness
- **Cloud Operations Model**:
  - Operating model definition
  - Support structure and escalation
  - Automation and tooling strategy
  - Cost management and optimization
- **Service Management**:
  - Service catalog and offerings
  - SLA and performance management
  - Capacity planning and scaling
  - Disaster recovery and business continuity

## (8) Financial Readiness
- **Financial Management**:
  - Cloud financial management maturity
  - Budgeting and forecasting processes
  - Chargeback/showback mechanisms
  - Cost allocation and tagging strategy
- **Investment & Funding**:
  - Migration budget and funding approval
  - Business case and ROI tracking
  - Cost optimization initiatives
  - Financial governance

## (9) Risk Assessment & Gap Analysis
- **Critical Risks**:
  - High-priority risks and blockers
  - Risk mitigation strategies
  - Dependencies and constraints
  - Timeline and resource risks
- **Readiness Gaps**:
  - Capability gaps by dimension
  - Gap severity and impact
  - Remediation priorities
  - Quick wins vs. long-term improvements

## (10) Recommendations & Action Plan
- **Strategic Recommendations**:
  - Prioritized improvement areas
  - Capability building roadmap
  - Partner and AWS engagement
  - Success metrics and KPIs
- **Action Plan**:
  - Short-term actions (0-3 months)
  - Medium-term actions (3-6 months)
  - Long-term actions (6-12 months)
  - Ownership and accountability

**IMPORTANT**: Base your analysis strictly on the content found in the MRA document. 
Do not make assumptions or add information not present in the assessment. Focus on extracting 
actionable insights that will inform the business case for AWS migration.

Format your response in markdown with clear headings, bullet points, and tables where appropriate.
"""

# Create the agent with MRA reading tools
agent = Agent(
    model=bedrock_model,
    system_prompt=system_message,
    tools=[read_docx_file, read_markdown_file, read_pdf_file]
)

# Example usage (commented out)
# question = """
# Analyze the Migration Readiness Assessment (MRA) document available in the input folder.
# Provide a comprehensive analysis covering all readiness dimensions and actionable recommendations
# for improving migration readiness and informing the business case.
# """
# 
# result = agent(question)
# print(result.message)


def extract_cloud_readiness(mra_content: str) -> int:
    """
    Extract cloud readiness score from MRA content (1-5 scale).
    
    Args:
        mra_content: MRA document text content
    
    Returns:
        Score from 1 (low) to 5 (high)
    """
    if not mra_content:
        return 3  # Default: Moderate
    
    content_lower = mra_content.lower()
    
    # Keywords indicating high readiness
    high_keywords = ['strong', 'excellent', 'mature', 'advanced', 'ready', 'prepared']
    # Keywords indicating low readiness
    low_keywords = ['weak', 'poor', 'immature', 'limited', 'lacking', 'unprepared']
    
    # Count keyword occurrences in cloud readiness context
    high_count = sum(1 for keyword in high_keywords if keyword in content_lower)
    low_count = sum(1 for keyword in low_keywords if keyword in content_lower)
    
    # Calculate score based on keyword balance
    if high_count > low_count + 2:
        return 5
    elif high_count > low_count:
        return 4
    elif low_count > high_count:
        return 2
    elif low_count > high_count + 2:
        return 1
    else:
        return 3


def extract_container_expertise(mra_content: str) -> int:
    """
    Extract container/Kubernetes expertise from MRA (1-5 scale).
    
    Args:
        mra_content: MRA document text content
    
    Returns:
        Score from 1 (none) to 5 (expert)
    """
    if not mra_content:
        return 3  # Default: Moderate
    
    content_lower = mra_content.lower()
    
    # Check for container/K8s mentions
    container_keywords = ['container', 'docker', 'kubernetes', 'k8s', 'eks', 'containerization']
    expertise_levels = {
        'expert': 5,
        'experienced': 4,
        'familiar': 3,
        'learning': 2,
        'none': 1,
        'no experience': 1
    }
    
    # Count container mentions
    container_mentions = sum(1 for keyword in container_keywords if keyword in content_lower)
    
    # Look for expertise level keywords
    for level, score in expertise_levels.items():
        if level in content_lower and any(k in content_lower for k in container_keywords):
            return score
    
    # If containers mentioned but no expertise level, assume moderate
    if container_mentions > 0:
        return 3
    
    # No container mentions = no expertise
    return 1


def extract_devops_maturity(mra_content: str) -> int:
    """
    Extract DevOps maturity from MRA (1-5 scale).
    
    Args:
        mra_content: MRA document text content
    
    Returns:
        Score from 1 (low) to 5 (high)
    """
    if not mra_content:
        return 3  # Default: Moderate
    
    content_lower = mra_content.lower()
    
    # DevOps indicators
    devops_keywords = ['ci/cd', 'continuous integration', 'continuous deployment', 
                      'automation', 'infrastructure as code', 'iac', 'terraform', 
                      'ansible', 'jenkins', 'gitlab', 'github actions']
    
    # Count DevOps practice mentions
    devops_count = sum(1 for keyword in devops_keywords if keyword in content_lower)
    
    # Score based on number of practices mentioned
    if devops_count >= 5:
        return 5  # High maturity
    elif devops_count >= 3:
        return 4  # Good maturity
    elif devops_count >= 2:
        return 3  # Moderate
    elif devops_count >= 1:
        return 2  # Developing
    else:
        return 1  # Low maturity


def extract_change_readiness(mra_content: str) -> int:
    """
    Extract change management readiness from MRA (1-5 scale).
    
    Args:
        mra_content: MRA document text content
    
    Returns:
        Score from 1 (low) to 5 (high)
    """
    if not mra_content:
        return 3  # Default: Moderate
    
    content_lower = mra_content.lower()
    
    # Change management indicators
    positive_keywords = ['change management', 'agile', 'flexible', 'adaptable', 'innovative']
    negative_keywords = ['resistant', 'rigid', 'traditional', 'risk-averse', 'slow']
    
    positive_count = sum(1 for keyword in positive_keywords if keyword in content_lower)
    negative_count = sum(1 for keyword in negative_keywords if keyword in content_lower)
    
    if positive_count > negative_count + 1:
        return 5
    elif positive_count > negative_count:
        return 4
    elif negative_count > positive_count:
        return 2
    elif negative_count > positive_count + 1:
        return 1
    else:
        return 3


def calculate_migration_timeline(mra_scores: dict, vm_categorization: dict) -> dict:
    """
    Calculate migration timeline based on MRA scores and VM distribution.
    
    Args:
        mra_scores: Dictionary with MRA scores
        vm_categorization: VM categorization results
    
    Returns:
        Dictionary with timeline breakdown
    """
    # Base timeline assumptions (weeks)
    base_timeline = {
        'assessment_weeks': 4,
        'pilot_weeks': 8,
        'days_per_simple_app': 3,
        'days_per_complex_app': 10
    }
    
    # Adjust based on MRA scores
    container_expertise = mra_scores.get('container_expertise_score', 3)
    devops_maturity = mra_scores.get('devops_maturity_score', 3)
    
    # Lower expertise = longer timeline
    if container_expertise < 3:
        base_timeline['pilot_weeks'] = int(base_timeline['pilot_weeks'] * 1.5)
        base_timeline['days_per_simple_app'] = int(base_timeline['days_per_simple_app'] * 1.3)
    
    if devops_maturity < 3:
        base_timeline['days_per_complex_app'] = int(base_timeline['days_per_complex_app'] * 1.5)
    
    # Calculate total timeline
    eks_vms = vm_categorization.get('eks_total', 0)
    
    # Estimate: 70% simple apps, 30% complex apps
    simple_apps = int(eks_vms * 0.70)
    complex_apps = int(eks_vms * 0.30)
    
    wave1_days = (simple_apps * base_timeline['days_per_simple_app'] + 
                  complex_apps * base_timeline['days_per_complex_app']) / 5  # 5 parallel teams
    
    total_weeks = (base_timeline['assessment_weeks'] + 
                   base_timeline['pilot_weeks'] + 
                   int(wave1_days / 5))  # Convert days to weeks
    
    return {
        'assessment_weeks': base_timeline['assessment_weeks'],
        'pilot_weeks': base_timeline['pilot_weeks'],
        'wave1_weeks': int(wave1_days / 5),
        'total_weeks': total_weeks,
        'total_months': int(total_weeks / 4)
    }


def parse_mra_for_eks(mra_filename: str = None) -> dict:
    """
    Parse MRA document and extract all scores for EKS recommendation.
    
    Args:
        mra_filename: MRA file name (auto-detects if None)
    
    Returns:
        Dictionary with all MRA scores
    """
    # Find MRA file if not specified
    if not mra_filename:
        mra_filename = find_mra_file()
    
    if not mra_filename:
        return {
            'cloud_readiness_score': 3,
            'container_expertise_score': 3,
            'devops_maturity_score': 3,
            'change_readiness_score': 3,
            'mra_available': False
        }
    
    # Read MRA content
    try:
        if mra_filename.endswith('.pdf'):
            content = read_pdf_file(mra_filename)
        elif mra_filename.endswith(('.docx', '.doc')):
            content = read_docx_file(mra_filename)
        elif mra_filename.endswith('.md'):
            content = read_markdown_file(mra_filename)
        else:
            content = ""
    except Exception as e:
        print(f"Warning: Could not read MRA file: {e}")
        content = ""
    
    # Extract scores
    return {
        'cloud_readiness_score': extract_cloud_readiness(content),
        'container_expertise_score': extract_container_expertise(content),
        'devops_maturity_score': extract_devops_maturity(content),
        'change_readiness_score': extract_change_readiness(content),
        'mra_available': True,
        'mra_filename': mra_filename
    }


if __name__ == "__main__":
    # Test the MRA analysis tools
    print("Testing Migration Readiness Assessment (MRA) Analysis Tools...")
    
    # Auto-detect MRA file
    mra_filename = find_mra_file()
    
    if not mra_filename:
        print("\n✗ No MRA file found in input directory")
        print("Expected files: mra-assessment.pdf, mra-assessment.docx, or mra-assessment.md")
    else:
        print(f"\n✓ Found MRA file: {mra_filename}")
        
        # Parse MRA for EKS
        mra_scores = parse_mra_for_eks(mra_filename)
        print(f"\n✓ MRA Scores:")
        print(f"  Cloud Readiness: {mra_scores['cloud_readiness_score']}/5")
        print(f"  Container Expertise: {mra_scores['container_expertise_score']}/5")
        print(f"  DevOps Maturity: {mra_scores['devops_maturity_score']}/5")
        print(f"  Change Readiness: {mra_scores['change_readiness_score']}/5")
